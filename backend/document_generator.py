from __future__ import annotations

from pathlib import Path
from uuid import uuid4

from fastapi import HTTPException
try:
    from docx import Document
except Exception:  # pragma: no cover - optional dependency
    Document = None


def crear_docx(secciones: dict, tmp_dir: Path) -> str:
    """Construye un archivo DOCX básico a partir de secciones."""
    if not Document:
        raise HTTPException(status_code=500, detail="Soporte DOCX no disponible")

    doc = Document()
    titulo = secciones.get("titulo", "Informe")
    doc.add_heading(titulo, level=1)

    intro = secciones.get("introduccion")
    if intro:
        doc.add_heading("Introducción", level=2)
        doc.add_paragraph(intro)

    desarrollo = secciones.get("desarrollo")
    if desarrollo:
        doc.add_heading("Desarrollo", level=2)
        for bloque in desarrollo.split("\n\n"):
            doc.add_paragraph(bloque)

    conclusion = secciones.get("conclusion")
    if conclusion:
        doc.add_heading("Conclusión", level=2)
        doc.add_paragraph(conclusion)

    path = tmp_dir / f"{uuid4()}.docx"
    doc.save(path)
    return str(path)
